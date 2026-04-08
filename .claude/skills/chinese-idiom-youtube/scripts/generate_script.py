"""
generate_script.py — 小雪的中文屋 script generator (trilingual: ZH / Pinyin / EN)

Each content block in sections.json must be a list of sentence objects:
  {"cn": "中文句子", "py": "zhōngwén jùzi", "en": "Chinese sentence"}

Usage:
    python generate_script.py \
        --idiom "过眼云烟" \
        --pinyin "guò yǎn yún yān" \
        --sections /tmp/sections.json \
        --output output/过眼云烟_Script.docx

sections.json schema:
{
  "opening_hook":    [{"cn":"...","py":"...","en":"..."}],
  "char_analyses": [
    {"char":"过","pinyin":"guò","lines":[{"cn":"...","py":"...","en":"..."}]},
    ...
  ],
  "combined_meaning": [{"cn":"...","py":"...","en":"..."}],
  "synonyms":         [{"cn":"...","py":"...","en":"..."}],
  "origin_meta": {"author":"北宋·苏轼","work":"《宝绘堂记》"},
  "origin_lines": [{"cn":"...","py":"...","en":"..."}],
  "origin_explanation": [{"cn":"...","py":"...","en":"..."}],
  "examples": [
    {"lines":[{"cn":"...","py":"...","en":"..."}]},
    ...
  ],
  "clip_revisit":  [{"cn":"...","py":"...","en":"..."}],
  "usage_note":    [{"cn":"...","py":"...","en":"..."}],
  "closing_hint":  [{"cn":"...","py":"...","en":"..."}]
}
"""

import argparse, json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colours ─────────────────────────────────────────────────────────────────
BLUE_DARK  = RGBColor(0x1A, 0x5C, 0x9A)
BLUE_MID   = RGBColor(0x2E, 0x86, 0xC1)
ORANGE     = RGBColor(0xE6, 0x7E, 0x22)
GREEN      = RGBColor(0x1E, 0x8B, 0x4C)
PURPLE     = RGBColor(0x6C, 0x3A, 0x9E)
GRAY       = RGBColor(0xAA, 0xAA, 0xAA)
GRAY_DARK  = RGBColor(0x55, 0x55, 0x55)

# ── Helpers ──────────────────────────────────────────────────────────────────
def _divider(doc):
    p    = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    for k, v in [('w:val','single'),('w:sz','4'),('w:space','1'),('w:color','CCCCCC')]:
        bot.set(qn(k), v)
    pBdr.append(bot)
    pPr.append(pBdr)

def _heading(doc, text):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = BLUE_DARK
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)

def _char_label(doc, char, pinyin):
    """Bold orange label for character analysis header."""
    p  = doc.add_paragraph()
    r  = p.add_run(f"【{char}  {pinyin}】")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = ORANGE
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)

def _trilingual(doc, lines, note=None):
    """
    Render a list of sentence dicts {"cn","py","en"} in the format:
        中文:   sentence
        拼音:   pīnyīn
        영어:   English

    Blank line between sentence groups.
    `note` is an optional string rendered in italics below the block.
    """
    if isinstance(lines, str):
        # Fallback: plain string
        p = doc.add_paragraph()
        p.add_run(lines).font.size = Pt(11)
        return

    for sentence in lines:
        cn = sentence.get("cn", "")
        py = sentence.get("py", "")
        py = py[:1].lower() + py[1:] if py else py  # 핀인 첫 글자 소문자
        en = sentence.get("en", "")

        # 中文
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        lbl = p.add_run("中文:  ")
        lbl.bold = True
        lbl.font.size = Pt(11)
        lbl.font.color.rgb = BLUE_MID
        body = p.add_run(cn)
        body.font.size = Pt(11)

        # 拼音
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        lbl2 = p2.add_run("拼音:  ")
        lbl2.bold = True
        lbl2.font.size = Pt(10)
        lbl2.font.color.rgb = PURPLE
        body2 = p2.add_run(py)
        body2.font.size = Pt(10)
        body2.italic = True
        body2.font.color.rgb = GRAY_DARK

        # 영어
        p3 = doc.add_paragraph()
        p3.paragraph_format.space_after = Pt(8)
        lbl3 = p3.add_run("영어:  ")
        lbl3.bold = True
        lbl3.font.size = Pt(10)
        lbl3.font.color.rgb = GREEN
        body3 = p3.add_run(en)
        body3.font.size = Pt(10)
        body3.font.color.rgb = GRAY_DARK

    if note:
        pn = doc.add_paragraph()
        pn.paragraph_format.space_after = Pt(4)
        rn = pn.add_run(note)
        rn.italic = True
        rn.font.size = Pt(10)
        rn.font.color.rgb = GRAY

# ── Fixed lines ──────────────────────────────────────────────────────────────
OPENER = [
    {"cn": "大家好，欢迎来到小雪的中文屋。我是小雪，Snowy。",
     "py": "Dàjiā hǎo, huānyíng lái dào xiǎo xuě de zhōngwén wū. Wǒ shì Xiǎo Xuě, Snowy.",
     "en": "Hello everyone, welcome to Snowy's Chinese House. I'm Xiao Xue, Snowy."}
]
CLIP_CUE = [
    {"cn": "那么今天就教大家一个成语。我们还是先看一段影视片段。",
     "py": "Nàme jīntiān jiù jiāo dàjiā yīgè chéngyǔ. Wǒmen háishi xiān kàn yīduàn yǐngshì piànduàn.",
     "en": "So today let's learn a Chinese idiom. Let's first watch a short drama clip."}
]
INTRO_CUE_TEMPLATE = lambda idiom: [
    {"cn": f"看完了片段，大家猜到今天的成语了吗？对，就是——{idiom}。{idiom}。{idiom}。",
     "py": "Kàn wán le piànduàn, dàjiā cāi dào jīntiān de chéngyǔ le ma? Duì, jiù shì——",
     "en": f"After watching, did everyone guess today's idiom? That's right — it's {idiom}. {idiom}. {idiom}."},
    {"cn": "我们还是先来分析一下每一个字的意思。",
     "py": "Wǒmen háishi xiān lái fēnxī yīxià měi yīgè zì de yìsi.",
     "en": "Let's start by analysing the meaning of each character."}
]
CLIP_REVISIT_CUE = [
    {"cn": "那我们回头再来看一下开头的影视片段——",
     "py": "Nà wǒmen huítóu zài lái kàn yīxià kāitóu de yǐngshì piànduàn——",
     "en": "Now let's go back and look at the opening drama clip again——"}
]
CLOSER = [
    {"cn": "好了，今天就到这里。希望你学到了一点新的中文。",
     "py": "Hǎo le, jīntiān jiù dào zhèlǐ. Xīwàng nǐ xué dào le yīdiǎn xīn de zhōngwén.",
     "en": "That's all for today. I hope you learned something new in Chinese."},
    {"cn": "如果你喜欢我的内容，欢迎点赞加关注小雪的中文屋。我们下次再见！",
     "py": "Rúguǒ nǐ xǐhuān wǒ de nèiróng, huānyíng diǎnzàn jiā guānzhù Xiǎo Xuě de Zhōngwén Wū. Wǒmen xià cì zàijiàn!",
     "en": "If you enjoy my content, please like and subscribe to Snowy's Chinese House. See you next time!"}
]

# ── Doc builder ──────────────────────────────────────────────────────────────
def build_doc(idiom, idiom_pinyin, s, output_path):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin    = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin   = Inches(1.2)
    sec.right_margin  = Inches(1.2)

    # ── Title block ────────────────────────────────────────────────────────
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = tp.add_run(idiom)
    t.font.size = Pt(28); t.bold = True; t.font.color.rgb = BLUE_DARK

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s1 = sp.add_run(idiom_pinyin)
    s1.font.size = Pt(13); s1.italic = True; s1.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c = cp.add_run("Chinese Idiom Series — 小雪的中文屋")
    c.font.size = Pt(10); c.font.color.rgb = GRAY

    doc.add_paragraph()
    _divider(doc)
    doc.add_paragraph()

    # ① Opening
    _heading(doc, "① 开场白")
    _trilingual(doc, OPENER)
    _trilingual(doc, s.get("opening_hook", []))
    _trilingual(doc, CLIP_CUE)
    p = doc.add_paragraph(); p.add_run("【影视片段】").italic = True

    # ② Idiom intro
    _heading(doc, "② 成语引出")
    _trilingual(doc, INTRO_CUE_TEMPLATE(idiom))

    # ③ Character analysis
    _heading(doc, "③ 字义分析")
    for ca in s.get("char_analyses", []):
        _char_label(doc, ca["char"], ca["pinyin"])
        _trilingual(doc, ca.get("lines", []))

    # ④ Combined meaning
    _heading(doc, "④ 成语整体含义")
    _trilingual(doc, s.get("combined_meaning", []))
    if s.get("synonyms"):
        _trilingual(doc, s["synonyms"])

    # ⑤ Origin
    _heading(doc, "⑤ 成语来源")
    meta = s.get("origin_meta", {})
    if meta:
        pm = doc.add_paragraph()
        rm = pm.add_run(f"{meta.get('author','')}  ·  {meta.get('work','')}")
        rm.font.size = Pt(11); rm.bold = True; rm.font.color.rgb = ORANGE
    _trilingual(doc, s.get("origin_lines", []))
    _trilingual(doc, s.get("origin_explanation", []))

    # ⑥ Examples + clip revisit
    _heading(doc, "⑥ 例句与片段回顾")
    for i, ex in enumerate(s.get("examples", []), 1):
        pm = doc.add_paragraph()
        rm = pm.add_run(f"例句 {i}")
        rm.bold = True; rm.font.size = Pt(11); rm.font.color.rgb = ORANGE
        _trilingual(doc, ex.get("lines", []))
    _trilingual(doc, CLIP_REVISIT_CUE)
    p = doc.add_paragraph(); p.add_run("【影视片段回放】").italic = True
    _trilingual(doc, s.get("clip_revisit", []))

    # ⑥-B Fill-in-the-blank quiz
    quiz = s.get("quiz", [])
    if quiz:
        _heading(doc, "⑥-B 填空练习")
        for i, q in enumerate(quiz, 1):
            pq = doc.add_paragraph()
            pq.paragraph_format.space_after = Pt(2)
            rn = pq.add_run(f"Q{i}.  ")
            rn.bold = True; rn.font.size = Pt(11); rn.font.color.rgb = BLUE_MID
            rq = pq.add_run(q.get("q_cn", ""))
            rq.font.size = Pt(12)
            if q.get("q_en"):
                ph = doc.add_paragraph()
                ph.paragraph_format.space_after = Pt(2)
                rh = ph.add_run("      " + q["q_en"])
                rh.italic = True; rh.font.size = Pt(10); rh.font.color.rgb = GRAY
            pa = doc.add_paragraph()
            pa.paragraph_format.space_after = Pt(10)
            pa.add_run("✅  ").font.size = Pt(11)
            ra = pa.add_run(q.get("full_cn", ""))
            ra.bold = True; ra.font.size = Pt(11); ra.font.color.rgb = GREEN
            if q.get("full_en"):
                pa2 = doc.add_paragraph()
                pa2.paragraph_format.space_after = Pt(12)
                ra2 = pa2.add_run("      " + q["full_en"])
                ra2.italic = True; ra2.font.size = Pt(10); ra2.font.color.rgb = GRAY

    # ⑦ Usage notes
    _heading(doc, "⑦ 使用注意")
    _trilingual(doc, s.get("usage_note", []))

    # ⑧ Closing
    _heading(doc, "⑧ 结尾")
    _trilingual(doc, s.get("closing_hint", []))
    _trilingual(doc, CLOSER)

    doc.add_paragraph()
    _divider(doc)
    np = doc.add_paragraph()
    np.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = np.add_run("✦ 小雪的中文屋 · Snowy's Chinese House ✦")
    nr.font.size = Pt(9); nr.font.color.rgb = GRAY

    doc.save(output_path)
    print(f"✅ Script saved: {output_path}")

if __name__ == "__main__":
    p = argparse.ArgumentParser()
    p.add_argument("--idiom",    required=True)
    p.add_argument("--pinyin",   required=True)
    p.add_argument("--sections", required=True)
    p.add_argument("--output",   required=True)
    args = p.parse_args()
    with open(args.sections, encoding="utf-8") as f:
        sections = json.load(f)
    build_doc(args.idiom, args.pinyin, sections, args.output)
